from __future__ import annotations

import json
import urllib.request
from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


PROJECT_DIR = Path(__file__).resolve().parents[1]
OUTPUT_PATH = PROJECT_DIR / "Описание карт Таро - review.docx"
API_BASE_URL = "http://localhost:8000"


DESCRIPTION_FIELDS = [
    ("1", "short_meaning", "Краткое значение"),
    ("2", "general", "Общее"),
    ("3", "love", "Любовь"),
    ("4", "career", "Работа"),
    ("5", "money", "Деньги"),
    ("6", "advice", "Совет"),
]


def fetch_json(path: str) -> list[dict]:
    with urllib.request.urlopen(f"{API_BASE_URL}{path}") as response:
        return json.loads(response.read().decode("utf-8"))


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shading = tc_pr.find(qn("w:shd"))
    if shading is None:
        shading = OxmlElement("w:shd")
        tc_pr.append(shading)
    shading.set(qn("w:fill"), fill)


def set_cell_text(cell, text: str, *, bold: bool = False, size: float = 8.5, color: str = "111111") -> None:
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = paragraph.add_run(text)
    run.bold = bold
    run.font.name = "Arial"
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_table_widths(table, widths: list[float]) -> None:
    for row in table.rows:
        for cell, width in zip(row.cells, widths):
            cell.width = Cm(width)


def configure_document(document: Document) -> None:
    section = document.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width = Cm(29.7)
    section.page_height = Cm(21.0)
    section.top_margin = Cm(1.2)
    section.bottom_margin = Cm(1.2)
    section.left_margin = Cm(1.2)
    section.right_margin = Cm(1.2)

    styles = document.styles
    styles["Normal"].font.name = "Arial"
    styles["Normal"].font.size = Pt(9)


def add_title(document: Document) -> None:
    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    title_run = title.add_run("Описание карт Таро - review")
    title_run.font.name = "Arial"
    title_run.font.size = Pt(20)
    title_run.font.bold = True

    subtitle = document.add_paragraph()
    subtitle.paragraph_format.space_after = Pt(8)
    subtitle_run = subtitle.add_run(
        "Таблица для ручной проверки текстов. У каждой карты шесть описаний: "
        "краткое значение, общее, любовь, работа, деньги, совет."
    )
    subtitle_run.font.name = "Arial"
    subtitle_run.font.size = Pt(9)
    subtitle_run.font.color.rgb = RGBColor(80, 80, 80)


def build_document() -> None:
    cards = fetch_json("/api/tarot-cards")
    interpretations = {
        item["card_code"]: item for item in fetch_json("/api/card-interpretations")
    }

    document = Document()
    configure_document(document)
    add_title(document)

    table = document.add_table(rows=1, cols=7)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    table.autofit = False

    headers = ["#", "Карта", "Код", "№", "Тип описания", "Текст", "Теги / score"]
    widths = [0.8, 3.0, 3.2, 0.7, 2.3, 13.2, 2.9]

    header_row = table.rows[0]
    set_repeat_table_header(header_row)
    for cell, header in zip(header_row.cells, headers):
        set_cell_shading(cell, "25212A")
        set_cell_text(cell, header, bold=True, size=8.5, color="FFFFFF")

    row_index = 1
    for card_index, card in enumerate(cards, start=1):
        interpretation = interpretations[card["card_code"]]
        tags = ", ".join(interpretation["tags_jsonb"])
        score = interpretation["score"]
        for number, field, label in DESCRIPTION_FIELDS:
            row = table.add_row()
            is_first = number == "1"
            values = [
                str(card_index) if is_first else "",
                card["name_ru"] if is_first else "",
                card["card_code"] if is_first else "",
                number,
                label,
                interpretation[field],
                f"{tags}\nscore: {score}" if is_first else "",
            ]
            for col_index, (cell, value) in enumerate(zip(row.cells, values)):
                set_cell_text(cell, value, size=8.2 if col_index == 5 else 8.0)
                if is_first:
                    set_cell_shading(cell, "F3EFE6")
                elif row_index % 2 == 0:
                    set_cell_shading(cell, "FAFAFA")
            row_index += 1

    set_table_widths(table, widths)
    document.save(OUTPUT_PATH)
    print(OUTPUT_PATH)


if __name__ == "__main__":
    build_document()
